'''
Backend script that does processes.
'''

import re
# import tk
from tkinter import filedialog
from docx import Document

def raise_error(error):
    match error:
        case "ERR_INV_FORMAT":
            return("Template is in invalid file format!")
        case "ERR_FILE_404":
            return("Template file not found at path!")
        case "ERR_NO_QUERIES":
            return("No queries found in the template!")
        case "ERR_USER_CANCEL":
            return("User cancelled the process.")
        case _:
            return("Unknown Error.")


# Restructures queries to send to UI
# in: list of queries
# out: list of 3 lists: [(original query,query name,(requested values))],[vars],[]
# (In case of dropdowns, requested values will be a tuple containing dictionary pairs of option and value)
def process_queries(queries)->list[list,list,list]:
    text_query_list = []
    text_query_name_set = [] #keeps track of query names to prevent duplication
    dropdown_list = [] #gets output
    dropdown_name_set = []
    calculation_list = []
    for query in queries:
        match query[2]:
            case 'Q':
                text_query_list.append(query)
            case 'C':
                calculation_list.append(query)
            case 'D':
                dropdown_list.append(query)
            case _:
                pass
    # Queries sorted by types.
    assembled_text_query_list = []
    assembled_dropdown_list = []
    # Get name and requested values:
    regex_tq = re.compile(r"{{Q:([^{},]+),?([^{}]*)}}")
    for raw_text_query in text_query_list:
        search_result = re.search(regex_tq, raw_text_query) #extract groups
        if search_result != None:
            query_name = search_result.group(1)
            query_description = search_result.group(2)
            #check if name exists in set
            if query_name not in text_query_name_set:
                if not query_description:
                    query_description = query_name
                text_query_name_set.append(query_name)
                assembled_query = [raw_text_query, query_name, query_description]
                assembled_text_query_list.append(assembled_query)
    
    regex_dq = re.compile(r"{{D:([^{},]+),?([^{}]*)}}")
    for raw_drop_query in dropdown_list:
        drop_search = re.search(regex_dq, raw_drop_query)
        if drop_search != None:
            drop_name = drop_search.group(1)
            drop_choices = drop_search.group(2)
            if drop_name not in dropdown_name_set:
                if drop_choices:
                    dropdown_name_set.append(drop_name)
                    assembled_dropdown = [raw_drop_query,drop_name,drop_choices]
                    assembled_dropdown_list.append(assembled_dropdown)

    regex_cl = re.compile(r"\[[^+-/*%!]*?\]")
    list_of_vars = []
    list_of_assembled_calc_queries = []
    for raw_calc_query in calculation_list:
        expression = raw_calc_query[4:-2]
        var_search = re.findall(regex_cl,raw_calc_query)
        if len(var_search) > 0:
            for var in var_search:
                assembled_calc_query = [raw_calc_query, expression]
                if var not in list_of_vars:
                    list_of_vars.append(var)
            list_of_assembled_calc_queries.append(assembled_calc_query)
    assembled_calc_list = [list_of_vars, list_of_assembled_calc_queries] #list of two lists: [all vars], [[raw, expression],..]

    return([assembled_text_query_list,assembled_calc_list,assembled_dropdown_list])
    
# Gets list of all queries [{{...}},]
def parse_txt(file_path):
    re_query_getter = re.compile(r"{{[^{}]*?}}")
    queries = []
    try:
        with open(file_path) as txt_file:
            for line in txt_file.readlines():
                queries_in_line = re.findall(re_query_getter,line)
                if len(queries_in_line) > 0:
                    for q in queries_in_line:
                        queries.append(q)
                
    except FileNotFoundError:
        return(raise_error("ERR_FILE_404"))
    
    if len(queries) > 0:
        return(process_queries(queries))
    else:
        return(raise_error("ERR_NO_QUERIES"))

def parse_docx(template_file):
    try:
        input_doc = Document(template_file)
        #get text and send to find queries
        re_query_getter = re.compile(r"{{[^{}]*?}}")
        queries = []
        for paragraph in input_doc.paragraphs:
            matches = re.findall(re_query_getter,paragraph.text)
            queries.extend(matches)
        for table in input_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(re_query_getter,paragraph.text)
                        queries.extend(matches)
        if len(queries) > 0:
            return process_queries(queries)
        else:
            return(raise_error("ERR_NO_QUERIES"))
    except:
        return(raise_error("ERR_UNKNOWN"))

def find_queries(template_file):
    file_ext = template_file[template_file.rindex('.'):]
    #print(file_ext)
    match file_ext:
        case '.txt':
            return(parse_txt(template_file))
        case '.docx':
           return(parse_docx(template_file))
        case _:
            return(raise_error("ERR_INV_FORMAT"))
# Called from ui, after file chosen.



def compile_txt(fq, rq, file_path):
    #text queries
    replacer_dict = {} # Original raw query:result
    for item in rq[0]:
        oq = item[0] #get original raw query
        rqn = item[1] #query name
        result = fq[0][rqn]
        replacer_dict[oq] = result
        # duplicate dict entries for repeated queries
        # tq_repeated_rx = re.compile(r"{{Q:[^,{}]+") #Finds text queries upto the first comma. add }} to get repeated query
        rtq = str("{{Q:"+str(rqn)+"}}")
        replacer_dict[rtq] = result
    # print(tq_replacer_dict)
    #dropdown queries
    for item in rq[2]:
        oq = item[0]
        oqn = item[1]
        result = fq[1][oqn]
        replacer_dict[oq] = result
        rtq = str("{{D:"+str(oqn)+"}}")
        replacer_dict[rtq] = result
    #calc queries TODO
    
    new_file_lines = []
    try:
        with open(file_path,"r", encoding="utf-8") as template:
            for line in template.readlines():
                new_line = line

                for key, value in replacer_dict.items():
                    if key in line:
                        new_line = new_line.replace(key,value)
                new_file_lines.append(new_line)
                    
        
        new_file_path = filedialog.asksaveasfilename(
            title="Save compiled file as...",
            defaultextension=".txt",
            filetypes=[("Text Files","*.txt"),("All Files","*.*")]
        )
        if new_file_path:
            try:
                with open(new_file_path,'w') as new_file:
                    print(f"Writing to {new_file_path}...")
                    # print(file_lines)
                    for l in new_file_lines:
                        new_file.write(f"{l}")
            except:
                return(raise_error("ERR_UNKNOWN"))
        else:
            return(raise_error("ERR_USER_CANCEL"))
    except:
        return(raise_error("ERR_UNKNOWN"))
    return("Document compiled successfully!")

def compile_docx(fq, rq, file_path):
    template = Document(file_path)
    #text queries
    replacer_dict = {} # Original raw query:result
    for item in rq[0]:
        oq = item[0] #get original raw query
        oqn = item[1] #query name
        result = fq[0][oqn]
        replacer_dict[oq] = result
        # duplicate dict entries for repeated queries
        # tq_repeated_rx = re.compile(r"{{Q:[^,{}]+") #Finds text queries upto the first comma. add }} to get repeated query
        rtq = str("{{Q:"+str(oqn)+"}}")
        replacer_dict[rtq] = result
    # print(tq_replacer_dict)
    #dropdown queries
    for item in rq[2]:
        oq = item[0]
        oqn = item[1]
        result = fq[1][oqn]
        replacer_dict[oq] = result
        rtq = str("{{D:"+str(oqn)+"}}")
        replacer_dict[rtq] = result

    #calc queries TODO

    # CODE FOR THIS FUNCTION WAS GENERATED BY GEMINI FLASH.
    def replace_in_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for target_tag, new_value in replacer_dict.items():
                while target_tag in paragraph.text:
                    # We look at the total combined text and find the character index of the tag
                    full_text = paragraph.text
                    start_idx = full_text.find(target_tag)
                    end_idx = start_idx + len(target_tag)
                    
                    current_idx = 0
                    start_run_idx = None
                    end_run_idx = None
                    
                    # Locate which runs contain the start and end of our tag
                    for i, run in enumerate(paragraph.runs):
                        run_len = len(run.text)
                        if start_run_idx is None and current_idx <= start_idx < current_idx + run_len:
                            start_run_idx = i
                        if current_idx <= end_idx <= current_idx + run_len:
                            end_run_idx = i
                            break
                        current_idx += run_len
                    
                    if start_run_idx is not None and end_run_idx is not None:
                        # Case 1: The entire tag is inside a single run (clean & easy)
                        if start_run_idx == end_run_idx:
                            run = paragraph.runs[start_run_idx]
                            run.text = run.text.replace(target_tag, new_value)
                        
                        # Case 2: The tag is split across multiple runs
                        else:
                            # Reconstruct text index relative to the starting run
                            run_start_text_idx = start_idx - (current_idx - sum(len(r.text) for r in paragraph.runs[end_run_idx:]))
                            
                            # Clean up the intermediate runs
                            for r_idx in range(start_run_idx + 1, end_run_idx):
                                paragraph.runs[r_idx].text = ""
                                
                            # Handle the end run fragment
                            end_run = paragraph.runs[end_run_idx]
                            end_tag_chars = end_idx - (full_text.find(end_run.text, start_idx) if full_text.find(end_run.text, start_idx) != -1 else 0)
                            # Fallback bound safe crop
                            end_run.text = end_run.text[end_tag_chars:] if end_tag_chars <= len(end_run.text) else ""
                            
                            # Place the new value safely into the starting run
                            start_run = paragraph.runs[start_run_idx]
                            # Cut off where the tag started, append new value
                            tag_pos_in_start_run = start_run.text.find(target_tag[:2]) # Peek ahead
                            if tag_pos_in_start_run != -1:
                                start_run.text = start_run.text[:tag_pos_in_start_run] + new_value
                            else:
                                start_run.text = start_run.text + new_value
                                
                    else:
                        # Edge case safety fallback
                        paragraph.text = paragraph.text.replace(target_tag, new_value, 1)
                        break

    
    # replace regular text first
    replace_in_paragraphs(template.paragraphs)
    # replace paragraphs in every cell in tables
    for table in template.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
    
    # DEBUG PRINT
    # for p in template.paragraphs:
    #     print(p.text)
    
    new_file_path = filedialog.asksaveasfilename(
            title="Save compiled file as...",
            defaultextension=".docx",
            filetypes=[("Document File","*.docx"),("All Files","*.*")]
        )
    if new_file_path:
        try:
            template.save(new_file_path)
        except:
            return(raise_error("ERR_UNKNOWN"))
    else:
        return(raise_error("ERR_USER_CANCEL"))
    return("Document compiled successfully!")
        

def compile(file_path, filled_queries, raw_queries):
    file_ext = file_path[file_path.rindex('.'):]
    match file_ext:
        case '.txt':
            compile_txt(filled_queries, raw_queries, file_path)
        case '.docx':
            compile_docx(filled_queries, raw_queries, file_path)
        case _:
            return(raise_error("ERR_INV_FORMAT"))